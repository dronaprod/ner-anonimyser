#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import logging
import math
import os
import random
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

# Module-level logger; configured in main() to write to log/log_<datetime>.log and console
logger = logging.getLogger("pii_pipeline")

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

# Local modules: gliner, presidio (optional), and optional Qwen NER via Ollama
from gliner_module import GLiNER
from presidio_module import AnalyzerEngine, NlpEngineProvider
try:
    from qwen_ollama_ner_module import detect_pii_with_qwen_ollama as _qwen_ner_detect
    from qwen_ollama_ner_module import detect_language as _detect_language
except Exception:  # pragma: no cover
    _qwen_ner_detect = None
    _detect_language = None

# GLiNER model IDs: English uses xlarge + gretelai + urchade; Arabic uses gretelai + urchade + arabic only (no Presidio, no xlarge)
GLINER_XLARGE_ID = "knowledgator/gliner-x-large"
GLINER_GRETELAI_ID = "gretelai/gretel-gliner-bi-large-v1.0"
GLINER_URCHADE_ID = "urchade/gliner_large-v2.1"
GLINER_ARABIC_ID = "NAMAa-Space/gliner_arabic-v2.1"
NER_NAME_XLARGE = "gliner_xlarge"
NER_NAME_GRETELAI = "gretelai_gliner_large"
NER_NAME_URCHADE = "urchade_gliner_large_2.1_og"
NER_NAME_ARABIC = "gliner_arabic"


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".xlsx", ".txt"}
GLINER_PII_LABELS = [

    # Identity
    "person",
    "name",
    "first name",
    "last name",
    "date of birth",

    # Government IDs
    "aadhaar number",
    "pan number",
    "gst number",
    "national id",
    "tax id",
    "certificate license number",

    # Healthcare
    "medical record number",
    "health plan beneficiary number",

    # Contact
    "email address",
    "phone number",

    # Address
    "street address",
    "address",
    "city",
    "state",
    "postcode",
    "country",

    # Network / Device
    "ipv4 address",
    "ipv6 address",
    "device identifier",
    "unique identifier",

    # Organization IDs
    "employee id",
    "customer id",

    # Financial
    "account number",
    "bank routing number",

    # Vehicle
    "license plate number",
    "vehicle identifier",

    # Biometric
    "biometric identifier"
]
GLINER_PII_LABELS_SET = set(GLINER_PII_LABELS)
PRESIDIO_TO_SHARED_LABEL = {

    # Identity
    "PERSON": "person",

    # Contact
    "EMAIL_ADDRESS": "email address",
    "PHONE_NUMBER": "phone number",

    # Location
    "LOCATION": "location",

    # Dates
    "DATE_TIME": "date",

    # Government IDs
    "US_SSN": "ssn",
    "US_PASSPORT": "passport number",
    "US_DRIVER_LICENSE": "driver license number",

    # Indian IDs
    "IN_AADHAAR": "aadhaar number",
    "IN_PAN": "pan number",
    "IN_PASSPORT": "passport number",
    "IN_VEHICLE_REGISTRATION": "vehicle registration number",

    # Financial
    "CREDIT_CARD": "credit card number",
    "IBAN_CODE": "bank account number",
    "US_BANK_NUMBER": "bank account number",

    # Network
    "IP_ADDRESS": "ip address",
    "MAC_ADDRESS": "mac address",

    # Internet
    "URL": "url",

    # Healthcare
    "MEDICAL_LICENSE": "medical license number",

    # Crypto
    "CRYPTO": "crypto wallet address",
}

LONG_NUMBER_PATTERN = re.compile(r"\b\d{6,}\b")
CARD_LIKE_PATTERN = re.compile(r"\b(?:\d[ -]?){12,20}\b")
UUID_PATTERN = re.compile(
    r"\b[0-9a-fA-F]{8}-"
    r"[0-9a-fA-F]{4}-"
    r"[0-9a-fA-F]{4}-"
    r"[0-9a-fA-F]{4}-"
    r"[0-9a-fA-F]{12}\b",
)
ALPHANUMERIC_PATTERN = re.compile(r"\b[A-Za-z0-9]{12,}\b")
BASE64_PATTERN = re.compile(r"\b[A-Za-z0-9+/]{20,}={0,2}\b")
HEX_32_PATTERN = re.compile(r"\b[a-fA-F0-9]{32}\b")
HEX_40_PATTERN = re.compile(r"\b[a-fA-F0-9]{40}\b")
HEX_64_PATTERN = re.compile(r"\b[a-fA-F0-9]{64}\b")
OBFUSCATED_EMAIL_PATTERN = re.compile(
    r"\b\w+\s*(?:at|\(at\))\s*\w+\s*(?:dot|\.)\s*\w+\b",
    re.IGNORECASE,
)

# Indian ID deterministic audit patterns (structure-only; no content logged)
AADHAAR_PATTERN = re.compile(r"\b\d{4}[\s-]?\d{4}[\s-]?\d{4}\b")
PAN_PATTERN = re.compile(r"\b[A-Z]{5}\d{4}[A-Z]\b")
GST_NUMBER_PATTERN = re.compile(r"\b\d{2}[A-Z]{5}\d{4}[A-Z]\d[A-Z]\d\b")
UDYAM_PATTERN = re.compile(r"\bUDYAM-[A-Z]{2}-\d{2}-\d{6,7}\b", re.IGNORECASE)
INDIAN_ID_PATTERNS: list[tuple[Any, str]] = [
    (AADHAAR_PATTERN, "aadhaar"),
    (PAN_PATTERN, "pan"),
    (GST_NUMBER_PATTERN, "gst_number"),
    (UDYAM_PATTERN, "udyam_number"),
]

# Section/label phrases that NERs sometimes wrongly tag as PII. Drop these to improve precision.
PII_LABEL_PHRASES: frozenset[str] = frozenset({
    "aadhaar number", "pan number", "pan card information name", "sample identity data",
    "gst information", "gstin", "business name", "father's name", "registration date",
    "business type", "pan card information", "aadhaar card information",
    "dob", "d.o.b.", "date of birth", "patient care coordination report",
    "coordination", "patient care", "care coordination", "coordination report",
    "metformin 500mg", "metformin 500", "identity data", "for testing purposes only",
})

# Document/report titles or generic phrases wrongly tagged as organization. Drop when label is organization.
DOCUMENT_TITLE_ORGANIZATION_PHRASES: frozenset[str] = frozenset({
    "patient care coordination report", "sample identity data", "identity data (for testing purposes only)",
    "patient care", "care coordination", "coordination report", "for testing purposes only",
    "bid insurance",
})

# Tokens from entropy audit that are common English words, not PII. Do not add as suspicious_token.
COMMON_NON_PII_WORDS: frozenset[str] = frozenset({
    "coordination", "registration", "information", "report", "patient", "care", "clinical",
    "medical", "treatment", "therapy", "diagnosis", "prescription", "medication", "dosage",
    "administration", "documentation", "confidential", "summary", "review", "assessment",
})

# Medication-like pattern: "Word 500mg" or "Word 500" (avoids Metformin 500mg tagged as person).
MEDICATION_LIKE_PATTERN = re.compile(
    r"^[A-Za-z][A-Za-z\s\-]+\d+\s*(?:mg|ml|g|mcg|iu|units?)?$",
    re.IGNORECASE,
)

_PRESIDIO_ENGINE: Any = None
_PRESIDIO_INIT_FAILED = False

LOG_DIR = Path(__file__).resolve().parent / "log"
REPO_ROOT = Path(__file__).resolve().parent.parent
MSPRISIDIO_PY = REPO_ROOT / "msprisidio" / "venv" / "bin" / "python"
MSPRISIDIO_SCRIPT = REPO_ROOT / "msprisidio" / "run_presidio_analyze.py"
PRESIDIO_SUBPROCESS_TIMEOUT = 60


def setup_logging() -> None:
    """Create log directory and log file log_<datetime>.log; log to file and console."""
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    log_file = LOG_DIR / f"log_{timestamp}.log"
    logger.setLevel(logging.DEBUG)
    logger.handlers.clear()
    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(logging.INFO)
    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s", datefmt="%Y-%m-%d %H:%M:%S")
    fh.setFormatter(fmt)
    ch.setFormatter(fmt)
    logger.addHandler(fh)
    logger.addHandler(ch)
    qwen_ner_logger = logging.getLogger("qwen_ollama_ner_module")
    qwen_ner_logger.setLevel(logging.DEBUG)
    qwen_ner_logger.addHandler(fh)
    logger.info("Logging to %s", log_file)


@dataclass
class PiiDetection:
    text: str
    label: str
    score: float


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="PII anonymization pipeline with Presidio + GLiNER + Qwen.",
    )
    parser.add_argument(
        "--input-dir",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "cache_105" / "files",
        help="Directory with source files (pdf/docx/xlsx/txt). Default: cache_105/files.",
    )
    parser.add_argument(
        "--num-files",
        type=int,
        default=10,
        help="How many files to process (default: 10).",
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=1600,
        help="Chunk size in characters.",
    )
    parser.add_argument(
        "--chunk-overlap",
        type=int,
        default=200,
        help="Chunk overlap in characters.",
    )
    parser.add_argument(
        "--gliner-model",
        type=str,
        default=os.environ.get("ARMOR_GLINER_MODEL", "knowledgator/gliner-x-large"),
        help="GLiNER model id (default: knowledgator/gliner-x-large or ARMOR_GLINER_MODEL; use urchade/gliner_medium-v2.1 for less RAM).",
    )
    parser.add_argument(
        "--gliner-threshold",
        type=float,
        default=0.28,
        help="GLiNER confidence threshold (default 0.28 for better recall; raise to 0.35+ to reduce false positives).",
    )
    parser.add_argument(
        "--presidio-threshold",
        type=float,
        default=0.35,
        help="Presidio confidence threshold.",
    )
    parser.add_argument(
        "--use-qwen-ner",
        action="store_true",
        default=True,
        help="Use Qwen (Ollama) for PII NER (default: True).",
    )
    parser.add_argument(
        "--no-qwen-ner",
        action="store_false",
        dest="use_qwen_ner",
        help="Disable Qwen NER (Ollama).",
    )
    parser.add_argument(
        "--qwen-ner-threshold",
        type=float,
        default=0.5,
        help="Qwen NER confidence threshold (default: 0.5).",
    )
    parser.add_argument(
        "--min-ner-agreement",
        type=int,
        default=1,
        choices=[1, 2, 3],
        help="Minimum number of NER detectors (Presidio, GLiNER, Qwen) that must agree on a PII to keep it. 1=keep if any detector finds it (default), 2=at least two agree, 3=all must agree.",
    )
    parser.add_argument(
        "--min-agreement-names",
        type=int,
        default=1,
        help="For person/name entities only: keep if at least this many detectors agree (default: 1, so names found by one detector are still anonymised).",
    )
    parser.add_argument(
        "--qwen-python",
        type=str,
        default=sys.executable,
        help="Python executable used to run qwen/run_qwen_sit.py.",
    )
    parser.add_argument(
        "--qwen-script",
        type=Path,
        default=Path(__file__).resolve().parent / "run_qwen_ollama.py",
        help="Script for anonymized values (default: run_qwen_ollama.py, Qwen 3.5 via Ollama). Use run_qwen_anonymize.py for HF Qwen2.5-1.5B, run_qwen_stub.py for stub.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("PII-Anonymisation/output"),
        help="Output directory.",
    )
    parser.add_argument(
        "--report-dir",
        type=Path,
        default=Path(__file__).resolve().parent / "db" / "reports",
        help="Directory to write report JSON (used by UI). Default: ./db/reports",
    )
    parser.add_argument(
        "--no-report",
        action="store_true",
        help="Do not write report JSON or update config (faster).",
    )
    parser.add_argument(
        "--progress-file",
        type=Path,
        default=None,
        help="Write live progress JSON here (stage, file, chunk, counts; no content). Used by UI.",
    )
    parser.add_argument(
        "--files",
        type=str,
        default=None,
        help="Comma-separated filenames to process (only these; must exist in input-dir). Default: use num-files.",
    )
    return parser.parse_args()


def _report_progress(progress_file: Path | None, **kwargs: Any) -> None:
    """Write progress dict to file (no PII/content). Atomic write."""
    if not progress_file:
        return
    try:
        payload = {"running": True, **kwargs}
        tmp = progress_file.with_suffix(progress_file.suffix + ".tmp")
        tmp.write_text(json.dumps(payload, ensure_ascii=True), encoding="utf-8")
        tmp.replace(progress_file)
    except Exception as e:
        logger.debug("Progress write failed: %s", e)


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)
    if suffix == ".xlsx":
        return _extract_xlsx(file_path)
    if suffix == ".txt":
        return file_path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {file_path}")


def _extract_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(file_path: Path) -> str:
    doc = Document(str(file_path))
    lines = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def _extract_xlsx(file_path: Path) -> str:
    wb = load_workbook(filename=str(file_path), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"[Sheet: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return []
    if chunk_overlap >= chunk_size:
        chunk_overlap = 0
    chunks: list[str] = []
    start = 0
    step = chunk_size - chunk_overlap
    while start < len(cleaned):
        chunks.append(cleaned[start : start + chunk_size])
        start += step
    return chunks


def _normalize_gliner_person_labels(detections: list[PiiDetection]) -> list[PiiDetection]:
    """When GLiNER returns 'last name' or 'first name' for a multi-word span, treat as 'person' for better agreement with Qwen/Presidio."""
    out: list[PiiDetection] = []
    for p in detections:
        label_lower = p.label.lower().strip()
        if label_lower in ("last name", "first name") and " " in p.text.strip():
            out.append(PiiDetection(text=p.text, label="person", score=max(p.score, 0.6)))
        else:
            out.append(p)
    return out


def detect_pii_with_gliner(
    model: GLiNER,
    text: str,
    threshold: float,
) -> list[PiiDetection]:
    # GLiNER: zero-shot NER; lower threshold = higher recall. "person" and full-name normalization improve agreement.
    predictions = model.predict_entities(text, GLINER_PII_LABELS, threshold=threshold)
    dedup: dict[tuple[str, str], PiiDetection] = {}
    for pred in predictions:
        value = str(pred.get("text", "")).strip()
        label = str(pred.get("label", "")).strip()
        score = float(pred.get("score", 0.0))
        if not value:
            continue
        key = (value.lower(), label.lower())
        current = dedup.get(key)
        if current is None or score > current.score:
            dedup[key] = PiiDetection(text=value, label=label, score=score)
    raw_list = sorted(dedup.values(), key=lambda x: (-len(x.text), x.text.lower()))
    raw_list = _normalize_gliner_person_labels(raw_list)
    raw_list = _dedup_detections(raw_list)
    filtered = _filter_gliner_label_phrases(raw_list)
    return _normalize_pii_types_by_pattern(filtered)


def _map_to_shared_label(raw_label: str, mapping: dict[str, str]) -> str | None:
    normalized = raw_label.strip().upper()
    mapped = mapping.get(normalized)
    if mapped:
        return mapped
    fallback = normalized.replace("_", " ").lower()
    return fallback if fallback in GLINER_PII_LABELS_SET else None


def _dedup_detections(items: list[PiiDetection]) -> list[PiiDetection]:
    dedup: dict[tuple[str, str], PiiDetection] = {}
    for item in items:
        if not item.text.strip() or not item.label.strip():
            continue
        key = (item.text.strip().lower(), item.label.strip().lower())
        current = dedup.get(key)
        if current is None or item.score > current.score:
            dedup[key] = PiiDetection(text=item.text.strip(), label=item.label.strip().lower(), score=item.score)
    return sorted(dedup.values(), key=lambda x: (-len(x.text), x.text.lower()))


def _normalize_pii_types_by_pattern(detections: list[PiiDetection]) -> list[PiiDetection]:
    """Override wrong NER labels when text matches Indian ID patterns to improve precision."""
    out: list[PiiDetection] = []
    for p in detections:
        text = p.text.strip()
        label = p.label.lower().strip()
        if AADHAAR_PATTERN.search(text) and label in ("date", "date of birth", "date_of_birth", "date time"):
            out.append(PiiDetection(text=p.text, label="aadhaar number", score=max(p.score, 0.85)))
            continue
        if PAN_PATTERN.search(text) and label == "organization":
            out.append(PiiDetection(text=p.text, label="pan number", score=max(p.score, 0.85)))
            continue
        if GST_NUMBER_PATTERN.search(text) and label == "organization":
            out.append(PiiDetection(text=p.text, label="gst number", score=max(p.score, 0.85)))
            continue
        digits_only = re.sub(r"[\s\-]", "", text)
        if len(digits_only) == 12 and digits_only.isdigit() and label == "organization":
            out.append(PiiDetection(text=p.text, label="aadhaar number", score=max(p.score, 0.85)))
            continue
        out.append(p)
    return out


def _filter_presidio_false_positives(detections: list[PiiDetection]) -> list[PiiDetection]:
    """Fix Presidio misclassifications: date→aadhaar when value is Aadhaar; drop/remap organization FPs."""
    result: list[PiiDetection] = []
    seen: set[tuple[str, str]] = set()
    for p in detections:
        value = p.text.strip()
        value_lower = value.lower()
        label = p.label.lower().strip()
        aadhaar_match = AADHAAR_PATTERN.search(value)
        if aadhaar_match and label in ("date", "date of birth", "date_time"):
            span = aadhaar_match.group(0).strip()
            key = (span.lower(), "aadhaar number")
            if key not in seen:
                seen.add(key)
                result.append(PiiDetection(text=span, label="aadhaar number", score=max(p.score, 0.85)))
            continue
        if label == "organization":
            if value_lower in PII_LABEL_PHRASES or value_lower in DOCUMENT_TITLE_ORGANIZATION_PHRASES:
                continue
            # State + pincode (e.g. "Maharashtra - 400001") often misclassified as organization; remap to address.
            if re.match(r"^[A-Za-z][A-Za-z\s]+-\s*\d{6}$", value.strip()):
                key = (value.lower(), "address")
                if key not in seen:
                    seen.add(key)
                    result.append(PiiDetection(text=p.text, label="address", score=max(p.score, 0.8)))
                continue
            if PAN_PATTERN.search(value):
                key = (value.lower(), "pan number")
                if key not in seen:
                    seen.add(key)
                    result.append(PiiDetection(text=p.text, label="pan number", score=max(p.score, 0.85)))
                continue
            if GST_NUMBER_PATTERN.search(value):
                key = (value.lower(), "gst number")
                if key not in seen:
                    seen.add(key)
                    result.append(PiiDetection(text=p.text, label="gst number", score=max(p.score, 0.85)))
                continue
            if aadhaar_match:
                span = aadhaar_match.group(0).strip()
                key = (span.lower(), "aadhaar number")
                if key not in seen:
                    seen.add(key)
                    result.append(PiiDetection(text=span, label="aadhaar number", score=max(p.score, 0.85)))
                continue
        key = (value.lower(), label)
        if key not in seen:
            seen.add(key)
            result.append(p)
    return sorted(result, key=lambda x: (-len(x.text), x.text.lower()))


def _filter_gliner_label_phrases(detections: list[PiiDetection]) -> list[PiiDetection]:
    """Drop GLiNER detections where the value is a form/section label (e.g. 'Aadhaar Number'), not actual PII."""
    return [p for p in detections if p.text.strip().lower() not in PII_LABEL_PHRASES]


def _is_likely_medication(text: str) -> bool:
    """True if text looks like a medication + dose (e.g. 'Metformin 500mg'), not a person name."""
    t = text.strip()
    if not t or len(t) > 80:
        return False
    return bool(MEDICATION_LIKE_PATTERN.match(t))


def _precision_filter(detections: list[PiiDetection]) -> list[PiiDetection]:
    """
    Drop common false positives across Presidio, GLiNER, Qwen to improve precision.
    - person/name: drop when text is purely numeric or looks like medication (e.g. Metformin 500mg).
    - organization: drop when text is a document/report title phrase.
    - date of birth: drop when text is only the label (e.g. 'DOB') not an actual date.
    - suspicious_token: drop when token is a common non-PII word (e.g. Coordination).
    - bank account number: drop when short numeric string (6–9 digits) likely internal ID.
    """
    result: list[PiiDetection] = []
    for p in detections:
        text = p.text.strip()
        value_lower = text.lower()
        label = p.label.lower().strip()
        if not text:
            continue
        if label in ("person", "name", "first name", "last name"):
            if text.isdigit():
                continue
            if _is_likely_medication(text):
                continue
        if label == "organization":
            if value_lower in DOCUMENT_TITLE_ORGANIZATION_PHRASES:
                continue
            if value_lower in PII_LABEL_PHRASES:
                continue
        if label in ("date of birth", "date_of_birth", "dob"):
            if value_lower in ("dob", "d.o.b.", "date of birth") and not re.search(r"\d", text):
                continue
        if label == "suspicious_token":
            if value_lower in COMMON_NON_PII_WORDS:
                continue
        if label in ("bank account number", "bank account"):
            digits_only = re.sub(r"\D", "", text)
            if 6 <= len(digits_only) <= 9 and digits_only.isdigit():
                continue
        result.append(p)
    return result


def shannon_entropy(text: str) -> float:
    if not text:
        return 0.0
    prob = [float(text.count(c)) / len(text) for c in set(text)]
    return -sum(p * math.log2(p) for p in prob)


def deterministic_audit_indian_ids(text: str) -> list[PiiDetection]:
    """Detect Indian IDs by structure only: Aadhaar, PAN, GST number, Udyam registration."""
    out: list[PiiDetection] = []
    seen: set[tuple[str, str]] = set()
    for pattern, label in INDIAN_ID_PATTERNS:
        for m in pattern.finditer(text):
            value = m.group(0).strip()
            if not value:
                continue
            key = (value.lower(), label)
            if key in seen:
                continue
            seen.add(key)
            out.append(PiiDetection(text=value, label=label, score=1.0))
    return sorted(out, key=lambda x: (-len(x.text), x.text.lower()))


def regex_entropy_audit(text: str) -> list[str]:
    findings: set[str] = set()
    patterns = [
        LONG_NUMBER_PATTERN,
        CARD_LIKE_PATTERN,
        UUID_PATTERN,
        ALPHANUMERIC_PATTERN,
        BASE64_PATTERN,
        HEX_32_PATTERN,
        HEX_40_PATTERN,
        HEX_64_PATTERN,
        OBFUSCATED_EMAIL_PATTERN,
    ]

    for pattern in patterns:
        for match in pattern.finditer(text):
            findings.add(match.group(0))

    for token in text.split():
        clean = token.strip(".,;:()[]{}<>\"'")
        if len(clean) >= 16:
            entropy = shannon_entropy(clean)
            if entropy >= 3.5:
                findings.add(clean)
    return list(findings)


def _ensure_spacy_model(model_name: str = "en_core_web_sm") -> None:
    """Ensure spaCy model is installed so Presidio can use it; download if missing."""
    try:
        import spacy
        spacy.load(model_name)
    except OSError:
        logger.info("Downloading spaCy model %s for Presidio...", model_name)
        try:
            import spacy.cli
            spacy.cli.download(model_name)
            import spacy
            spacy.load(model_name)
        except Exception as e:
            logger.warning("Could not auto-download spaCy model %s: %s", model_name, e)
            raise


def _get_msprisidio_site_packages() -> Path | None:
    """Return msprisidio venv site-packages path for current Python version if it exists."""
    py_ver = f"python{sys.version_info.major}.{sys.version_info.minor}"
    sp = REPO_ROOT / "msprisidio" / "venv" / "lib" / py_ver / "site-packages"
    return sp if sp.is_dir() else None


def _detect_pii_with_presidio_subprocess(text: str, threshold: float) -> list[PiiDetection]:
    """Run Presidio via msprisidio (venv python or current python + PYTHONPATH) when in-process Presidio is not available."""
    run_cmd: list[str] | None = None
    env = os.environ.copy()
    if MSPRISIDIO_PY.exists() and os.access(str(MSPRISIDIO_PY), os.X_OK) and MSPRISIDIO_SCRIPT.is_file():
        run_cmd = [str(MSPRISIDIO_PY), str(MSPRISIDIO_SCRIPT)]
    else:
        sp = _get_msprisidio_site_packages()
        if sp is not None and MSPRISIDIO_SCRIPT.is_file():
            env["PYTHONPATH"] = str(sp) + os.pathsep + env.get("PYTHONPATH", "")
            run_cmd = [sys.executable, str(MSPRISIDIO_SCRIPT)]
    if run_cmd is None:
        logger.debug(
            "Presidio: msprisidio venv not runnable and no msprisidio site-packages for this Python; install presidio-analyzer in this env or use same Python as msprisidio.",
        )
        return []
    try:
        proc = subprocess.run(
            run_cmd,
            input=text.encode("utf-8"),
            capture_output=True,
            timeout=PRESIDIO_SUBPROCESS_TIMEOUT,
            cwd=str(REPO_ROOT),
            env=env,
        )
        out = proc.stdout.decode("utf-8", errors="replace").strip()
        if proc.returncode != 0 and proc.stderr:
            logger.debug("Presidio subprocess stderr: %s", proc.stderr.decode("utf-8", errors="replace")[:500])
        if not out:
            return []
        raw = json.loads(out)
        if not isinstance(raw, list):
            return []
    except (subprocess.TimeoutExpired, json.JSONDecodeError, OSError) as e:
        logger.warning("Presidio subprocess failed: %s", e)
        return []
    dedup: dict[tuple[str, str], PiiDetection] = {}
    for r in raw:
        score = float(r.get("score", 0.0) or 0.0)
        if score < threshold:
            continue
        start = int(r.get("start", 0))
        end = int(r.get("end", 0))
        if end <= start:
            continue
        value = text[start:end].strip()
        if not value:
            continue
        raw_label = str(r.get("entity_type", "unknown")).strip()
        label = _map_to_shared_label(raw_label, PRESIDIO_TO_SHARED_LABEL)
        if not label:
            label = raw_label.lower().replace("_", " ").strip() or "entity"
        key = (value.lower(), label.lower())
        current = dedup.get(key)
        if current is None or score > current.score:
            dedup[key] = PiiDetection(text=value, label=label, score=score)
    raw_list = sorted(dedup.values(), key=lambda x: (-len(x.text), x.text.lower()))
    return _normalize_pii_types_by_pattern(_filter_presidio_false_positives(raw_list))


def detect_pii_with_presidio(text: str, threshold: float) -> list[PiiDetection]:
    global _PRESIDIO_ENGINE, _PRESIDIO_INIT_FAILED
    if AnalyzerEngine is None or _PRESIDIO_INIT_FAILED:
        return _detect_pii_with_presidio_subprocess(text, threshold)
    if _PRESIDIO_ENGINE is None:
        try:
            _ensure_spacy_model("en_core_web_sm")
            nlp_provider = NlpEngineProvider(
                nlp_configuration={
                    "nlp_engine_name": "spacy",
                    "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}],
                },
            )
            nlp_engine = nlp_provider.create_engine()
            _PRESIDIO_ENGINE = AnalyzerEngine(
                nlp_engine=nlp_engine,
                supported_languages=["en"],
            )
        except Exception as exc:  # pragma: no cover
            try:
                _PRESIDIO_ENGINE = AnalyzerEngine(supported_languages=["en"])
                logger.debug("Presidio engine created (default config)")
            except Exception as exc2:
                _PRESIDIO_INIT_FAILED = True
                logger.warning("Presidio unavailable, continuing without presidio detection: %s", exc2)
                return []
        else:
            logger.debug("Presidio engine created successfully")
    try:
        results = _PRESIDIO_ENGINE.analyze(text=text, language="en")
    except Exception as exc:  # pragma: no cover
        logger.warning("Presidio analyze failed for chunk, continuing: %s", exc)
        return []
    logger.debug("Presidio analyze returned %d raw results", len(results))
    dedup: dict[tuple[str, str], PiiDetection] = {}
    for result in results:
        score = float(getattr(result, "score", 0.0) or 0.0)
        if score < threshold:
            continue
        start = int(getattr(result, "start", 0))
        end = int(getattr(result, "end", 0))
        if end <= start:
            continue
        value = text[start:end].strip()
        if not value:
            continue
        raw_label = str(getattr(result, "entity_type", "unknown")).strip()
        label = _map_to_shared_label(raw_label, PRESIDIO_TO_SHARED_LABEL)
        if not label:
            label = raw_label.lower().replace("_", " ").strip() or "entity"
        key = (value.lower(), label.lower())
        current = dedup.get(key)
        if current is None or score > current.score:
            dedup[key] = PiiDetection(text=value, label=label, score=score)
    raw_list = sorted(dedup.values(), key=lambda x: (-len(x.text), x.text.lower()))
    return _normalize_pii_types_by_pattern(_filter_presidio_false_positives(raw_list))


def _canonical_pii_label(label: str) -> str:
    """Return canonical form for PII label so aadhaar/aadhaar number and pan/pan number merge."""
    L = label.lower().strip()
    if L in ("aadhaar", "aadhaar number"):
        return "aadhaar number"
    if L in ("pan", "pan number"):
        return "pan number"
    if L in ("gst_number", "gst number"):
        return "gst number"
    if L in ("date_of_birth", "date of birth", "dob"):
        return "date of birth"
    return L


def _label_agreement(l1: str, l2: str) -> bool:
    """True if two labels are considered the same for ensemble agreement and dedupe (e.g. person/name, email/email address, date/date of birth)."""
    a, b = l1.lower().strip(), l2.lower().strip()
    if a == b:
        return True
    if _canonical_pii_label(a) == _canonical_pii_label(b):
        return True
    if a in ("person", "name") and b in ("person", "name"):
        return True
    if a in ("email", "email address") and b in ("email", "email address"):
        return True
    if "organization" in a and "organization" in b:
        return True
    if "address" in a or "address" in b:
        if a in ("state", "location", "city", "address", "street address") or "address" in a:
            if b in ("state", "location", "city", "address", "street address") or "address" in b:
                return True
    if a in ("state", "location", "city", "street address") and b in ("state", "location", "city", "address", "street address"):
        return True
    if a in ("date", "date of birth", "date_of_birth") and b in ("date", "date of birth", "date_of_birth"):
        return True
    if a in ("phone number", "phone") and b in ("phone number", "phone"):
        return True
    return False


def _group_contains_detection(group: list[PiiDetection], text: str, label: str) -> bool:
    t = text.lower()
    for p in group:
        if p.text.lower() == t and _label_agreement(p.label, label):
            return True
    return False


# Lazy-loaded GLiNER models (model_id -> GLiNER) to avoid loading all 4 at once
_gliner_model_cache: dict[str, Any] = {}

def _get_gliner_model(model_id: str) -> GLiNER:
    if model_id not in _gliner_model_cache:
        logger.info("Loading GLiNER model: %s", model_id)
        _gliner_model_cache[model_id] = GLiNER.from_pretrained(model_id)
    return _gliner_model_cache[model_id]


def _dedupe_findings(findings: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """
    Group duplicate or semi-duplicate findings into one row; union found_by and take best score/type.
    (1) Same value (case-insensitive): always merge into one (regardless of type); use type with highest score, union found_by.
    (2) Contained/containing values with compatible label (e.g. John vs John Smith): merge into one row.
    """
    if not findings:
        return []
    # Normalize: value_lower, pii_type, found_by set, score, value_orig
    normalized: list[tuple[str, str, set[str], float, str]] = []
    for f in findings:
        value = (f.get("value") or "").strip()
        pii_type = (f.get("pii_type") or "").strip()
        if not value:
            continue
        if not pii_type:
            pii_type = "entity"
        canonical = _canonical_pii_label(pii_type)
        found_by = f.get("found_by") or []
        by_set = set(str(x).strip().lower() for x in found_by if x) if isinstance(found_by, list) else set()
        score = float(f.get("score", 0.8))
        normalized.append((value.lower(), canonical, by_set, score, value))

    # Pass 1: group by exact same value (case-insensitive); merge all into one per value
    by_value: dict[str, list[tuple[str, str, set[str], float, str]]] = {}
    for n in normalized:
        v_lower = n[0]
        if v_lower not in by_value:
            by_value[v_lower] = []
        by_value[v_lower].append(n)

    pass1: list[tuple[str, str, set[str], float, str]] = []
    for _v_lower, group in by_value.items():
        by_union: set[str] = set()
        best_score = 0.0
        best_orig = ""
        best_type = ""
        for _v_l, can, by_set, score, v_orig in group:
            by_union |= by_set
            if score > best_score or (score == best_score and len(v_orig) > len(best_orig)):
                best_score = score
                best_orig = v_orig
                best_type = can
        if not best_orig:
            best_orig = group[0][4]
            best_type = group[0][1]
            best_score = group[0][3]
        pass1.append((_v_lower, best_type, by_union, best_score, best_orig))

    # Pass 2: merge contained/containing with compatible label (e.g. John and John Smith)
    groups: list[list[tuple[str, str, set[str], float, str]]] = []
    for n in pass1:
        v_lower, can, by_set, score, v_orig = n
        merged = False
        for g in groups:
            if not _label_agreement(g[0][1], can):
                continue
            for _v_l, _c, _by, _sc, _orig in g:
                if v_lower == _v_l:
                    g.append(n)
                    merged = True
                    break
                if v_lower in _v_l or _v_l in v_lower:
                    g.append(n)
                    merged = True
                    break
            if merged:
                break
        if not merged:
            groups.append([n])

    out: list[dict[str, Any]] = []
    for g in groups:
        by_union = set()
        best_val = ""
        best_score = 0.0
        best_type = ""
        for _v_l, _c, _by, _sc, _orig in g:
            by_union |= _by
            if len(_orig) > len(best_val) or (len(_orig) == len(best_val) and _sc > best_score):
                best_val = _orig
                best_score = _sc
                best_type = _c
        if not best_type:
            best_type = g[0][1]
        out.append({
            "value": best_val,
            "pii_type": best_type,
            "score": round(best_score, 4),
            "found_by": sorted(by_union),
        })
    return out


def merge_pii_detections(
    detection_groups: list[list[PiiDetection]],
) -> list[PiiDetection]:
    merged: dict[tuple[str, str], PiiDetection] = {}
    for group in detection_groups:
        for item in group:
            canonical = _canonical_pii_label(item.label)
            key = (item.text.lower(), canonical)
            existing = merged.get(key)
            if existing is None or item.score > existing.score:
                merged[key] = PiiDetection(text=item.text, label=canonical, score=item.score)
    return sorted(merged.values(), key=lambda x: (-len(x.text), x.text.lower()))


def _is_person_name_label(label: str) -> bool:
    """True if label is person/name for relaxed agreement."""
    L = label.lower().strip()
    return L in ("person", "name")


def pii_ensemble_agreement(
    detection_groups: list[list[PiiDetection]],
    min_agreement: int,
    min_agreement_for_names: int = 1,
) -> list[PiiDetection]:
    """
    Keep only PIIs that at least min_agreement detectors found (ensemble agreement).
    For person/name entities, use min_agreement_for_names (default 1) so names
    found by only one detector are still anonymised.
    """
    union = merge_pii_detections(detection_groups)
    if min_agreement <= 1 and min_agreement_for_names <= 1:
        return union
    agreed: list[PiiDetection] = []
    for p in union:
        count = sum(
            1 for group in detection_groups if _group_contains_detection(group, p.text, p.label)
        )
        required = min_agreement_for_names if _is_person_name_label(p.label) else min_agreement
        if count >= required:
            agreed.append(p)
    return sorted(agreed, key=lambda x: (-len(x.text), x.text.lower()))


def call_qwen_json(
    qwen_python: str,
    qwen_script: Path,
    system_prompt: str,
    user_prompt: str,
) -> dict[str, Any]:
    if not qwen_script.is_file():
        raise FileNotFoundError(f"Qwen script not found: {qwen_script}")

    payload = json.dumps({"system_prompt": system_prompt, "user_prompt": user_prompt}) + "\n"
    completed = subprocess.run(
        [qwen_python, str(qwen_script)],
        input=payload,
        text=True,
        capture_output=True,
        check=False,
    )
    if completed.returncode != 0:
        raise RuntimeError(
            f"Qwen invocation failed with code {completed.returncode}: {completed.stderr.strip()}",
        )

    output = completed.stdout.strip()
    if not output:
        return {"error": "Empty response from Qwen"}
    return parse_json_like(output)


def parse_json_like(text: str) -> dict[str, Any]:
    text = text.strip()
    if not text:
        return {"error": "Empty text"}

    candidates = [text]
    fenced = re.search(r"```(?:json)?\s*(\{.*?\}|\[.*?\])\s*```", text, re.DOTALL)
    if fenced:
        candidates.append(fenced.group(1).strip())
    for start_match in re.finditer(r"[\{\[]", text):
        snippet = text[start_match.start() :]
        end_obj = snippet.rfind("}")
        end_arr = snippet.rfind("]")
        end = max(end_obj, end_arr)
        if end >= 0:
            candidates.append(snippet[: end + 1].strip())

    for candidate in candidates:
        try:
            loaded = json.loads(candidate)
            if isinstance(loaded, dict):
                return loaded
            return {"result": loaded}
        except json.JSONDecodeError:
            continue
    return {"raw_response": text}


def build_qwen_replacement_prompt(
    chunk: str, detected_pii: list[PiiDetection], lang: str = "en"
) -> tuple[str, str]:
    lang_rule = (
        "**Language**: The text is in Arabic. Generate ALL anonymized replacement values in Arabic only "
        "(e.g. Arabic names, Arabic addresses, Arabic organisation names). Do not use English words for replacements."
        if lang == "ar"
        else "**Language**: The text is in English (or another Latin-script language). Generate ALL anonymized replacement values in English only."
    )
    system = (
        "You are a PII anonymization assistant. Return JSON only: "
        "{\"replacements\":[{\"original_value\":\"...\",\"anonymized_value\":\"...\",\"pii_type\":\"...\"}]}. "
        "Rules: "
        "(1) " + lang_rule + " "
        "(2) **Same type**: name→name, date→date, phone→phone, email→email, organisation→organisation, etc. "
        "(3) **Structurally and contextually similar**: "
        "Dates: preserve the exact format (DD/MM/YYYY vs MM/DD/YYYY vs YYYY-MM-DD, month names, separators). Same era/century if obvious. "
        "Phones: preserve country code pattern and separators (e.g. +46..., 0xx..., (0xx) ...). "
        "IDs/numbers: preserve length and separator pattern (e.g. SSN dashes, card spaces). "
        "Addresses: same country/region style (street format, postal pattern). "
        "Output only the JSON object."
    )
    detected_payload = [
        {"value": item.text, "pii_type": item.label, "confidence": round(item.score, 4)}
        for item in detected_pii
    ]
    user = (
        "PII detected (Presidio + GLiNER + Qwen):\n"
        f"{json.dumps(detected_payload, ensure_ascii=False)}\n\n"
        "Text chunk:\n"
        f"{chunk}\n\n"
        "Generate one replacement per PII. Use the SAME language as the text for all replacement values (Arabic→Arabic, English→English). Same type and structure. JSON only."
    )
    return system, user


def normalize_label(label: str) -> str:
    return re.sub(r"\s+", "_", label.strip().lower())


def partial_mask(text: str, visible_ratio: float = 0.3) -> str:
    """Mask text with #, leaving approximately `visible_ratio` of characters visible (at start and end)."""
    if not text or not text.strip():
        return "#####"
    s = text.strip()
    n = len(s)
    if n <= 2:
        return "#" * n
    keep_each = max(1, round(n * (visible_ratio / 2)))
    if keep_each * 2 >= n:
        keep_each = max(1, n // 3)
    mid = n - 2 * keep_each
    return s[:keep_each] + ("#" * mid) + s[-keep_each:]


def _infer_date_format(original: str) -> str:
    """Infer date format from original string: ddmmyyyy, mmddyyyy, yyyymmdd, or default mmddyyyy."""
    s = original.strip()
    # YYYY-MM-DD or YYYY/MM/DD
    if re.search(r"\b(19|20)\d{2}[-/]\d{1,2}[-/]\d{1,2}\b", s):
        return "yyyymmdd"
    # DD/MM/YYYY or DD-MM-YYYY (day typically > 12 when ambiguous)
    if re.search(r"\b\d{1,2}[-/]\d{1,2}[-/](19|20)\d{2}\b", s):
        return "ddmmyyyy"
    # MM/DD/YYYY (US)
    if re.search(r"\b\d{1,2}[-/]\d{1,2}[-/](19|20)?\d{2}\b", s):
        return "mmddyyyy"
    # Month name (e.g. 15 March 2025, March 15, 2025)
    if re.search(r"[A-Za-z]+\s+\d{1,2},?\s+(19|20)\d{2}", s) or re.search(r"\d{1,2}\s+[A-Za-z]+\s+(19|20)\d{2}", s):
        return "mmddyyyy"  # arbitrary; hard to distinguish without locale
    return "mmddyyyy"


def _infer_phone_prefix(original: str) -> str:
    """Infer country/region prefix from original (e.g. +46, +1, 0). Return prefix for synthetic."""
    digits = re.sub(r"\D", "", original)
    s = original.strip()
    if s.startswith("0") and len(digits) >= 9:
        return "0"
    if s.startswith("+") and len(digits) >= 10:
        # Common country codes: 1, 33, 44, 46, 49, etc.
        if digits.startswith("1") and len(digits) >= 10:
            return "+1"
        if digits.startswith("44"):
            return "+44"
        if digits.startswith("46"):
            return "+46"
        if digits.startswith("33"):
            return "+33"
        if digits.startswith("49"):
            return "+49"
        if len(digits) >= 10:
            return "+" + digits[: min(3, 1 + (1 if digits[0] == "1" else 2))]
    return "+1"


def synthetic_value_for_type(pii_type: str, original: str, lang: str = "en") -> str:
    t = pii_type.lower()
    if "suspicious_token" in t:
        return partial_mask(original, visible_ratio=0.3)
    seed = abs(hash((pii_type.lower(), original))) % 1_000_000
    rng = random.Random(seed)
    if "email" in t:
        return f"user{seed % 100000}@example.com"
    if "phone" in t:
        prefix = _infer_phone_prefix(original)
        if prefix == "0":
            return f"0{rng.randint(70, 79)}{rng.randint(1000000, 9999999)}"
        if prefix.startswith("+46"):
            return f"+46{rng.randint(70, 79)}{rng.randint(100000, 999999)}"
        if prefix.startswith("+44"):
            return f"+44{rng.randint(7700, 7799)}{rng.randint(100000, 999999)}"
        if prefix.startswith("+33"):
            return f"+33{rng.randint(1, 9)}{rng.randint(10000000, 99999999)}"
        if prefix.startswith("+49"):
            return f"+49{rng.randint(150, 179)}{rng.randint(1000000, 9999999)}"
        return f"+1-555-{rng.randint(100, 999)}-{rng.randint(1000, 9999)}"
    if "name" in t or "person" in t:
        if lang == "ar":
            first = ["أحمد", "محمد", "علي", "خالد", "عمر", "يوسف"][seed % 6]
            last = ["العلي", "الحسن", "المرعي", "الشامي", "الغربي", "النجار"][(seed // 7) % 6]
            return f"{first} {last}"
        first = ["Alex", "Jordan", "Taylor", "Casey", "Morgan", "Avery"][seed % 6]
        last = ["Smith", "Johnson", "Clark", "Davis", "Miller", "Brown"][(seed // 7) % 6]
        return f"{first} {last}"
    if "address" in t:
        if lang == "ar":
            return f"شارع {rng.randint(10, 999)}، حي النور، الرياض"
        return f"{rng.randint(10, 999)} Example Street, Springfield"
    if "organization" in t or "company" in t:
        if lang == "ar":
            return f"شركة الخدمات المتحدة {seed % 1000}"
        return f"Acme Holdings {seed % 1000}"
    if "ssn" in t:
        return f"{rng.randint(100, 999)}-{rng.randint(10, 99)}-{rng.randint(1000, 9999)}"
    if "passport" in t:
        return f"P{rng.randint(10000000, 99999999)}"
    if "credit" in t and "card" in t:
        return f"{rng.randint(1000, 9999)} {rng.randint(1000, 9999)} {rng.randint(1000, 9999)} {rng.randint(1000, 9999)}"
    if "bank" in t and "account" in t:
        return "".join(str(rng.randint(0, 9)) for _ in range(12))
    if "ip" in t:
        return f"10.{rng.randint(0, 255)}.{rng.randint(0, 255)}.{rng.randint(1, 254)}"
    if "date" in t:
        fmt = _infer_date_format(original)
        y = rng.randint(1970, 1999)
        m = rng.randint(1, 12)
        d = rng.randint(1, 28)
        if fmt == "yyyymmdd":
            return f"{y}-{m:02d}-{d:02d}"
        if fmt == "ddmmyyyy":
            return f"{d:02d}/{m:02d}/{y}"
        return f"{m:02d}/{d:02d}/{y}"
    if "username" in t:
        return f"user_{seed % 100000}"
    if "aadhaar" in t:
        return f"{rng.randint(1000, 9999)} {rng.randint(1000, 9999)} {rng.randint(1000, 9999)}"
    if "pan" in t:
        return f"{chr(65 + seed % 26)}{chr(65 + (seed // 26) % 26)}{chr(65 + (seed // 676) % 26)}{chr(65 + (seed // 17576) % 26)}{chr(65 + (seed // 456976) % 26)}{rng.randint(1000, 9999)}{chr(65 + rng.randint(0, 25))}"
    if "gst_number" in t:
        return f"{rng.randint(10, 99)}{chr(65 + seed % 26)}{chr(65 + (seed // 26) % 26)}{chr(65 + (seed // 676) % 26)}{chr(65 + (seed // 17576) % 26)}{chr(65 + (seed // 456976) % 26)}{rng.randint(1000, 9999)}{chr(65 + rng.randint(0, 25))}{rng.randint(1, 9)}{chr(65 + rng.randint(0, 25))}{rng.randint(1, 9)}"
    if "udyam" in t:
        return f"UDYAM-{chr(65 + seed % 26)}{chr(65 + (seed // 26) % 26)}-{rng.randint(10, 99)}-{rng.randint(100000, 9999999)}"
    return f"<{normalize_label(pii_type)}_{seed % 100000}>"


def datatype_match(value: str, pii_type: str) -> bool:
    t = pii_type.lower()
    if "suspicious_token" in t:
        return False
    if "email" in t:
        return bool(re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", value))
    if "phone" in t:
        return bool(re.search(r"\d{7,}", re.sub(r"\D", "", value)))
    if "ssn" in t:
        return bool(re.fullmatch(r"\d{3}-\d{2}-\d{4}", value))
    if "ip" in t:
        return bool(re.fullmatch(r"(?:\d{1,3}\.){3}\d{1,3}", value))
    if "name" in t or "person" in t:
        return any(c.isalpha() for c in value) and len(value.strip()) >= 2
    if "credit" in t and "card" in t:
        digits = re.sub(r"\D", "", value)
        return 13 <= len(digits) <= 19
    if "bank" in t and "account" in t:
        digits = re.sub(r"\D", "", value)
        return len(digits) >= 8
    if "date" in t:
        return bool(re.search(r"\d", value))
    if "aadhaar" in t:
        return bool(AADHAAR_PATTERN.fullmatch(value.strip()))
    if "pan" in t:
        return bool(PAN_PATTERN.fullmatch(value.strip()))
    if "gst_number" in t:
        return bool(GST_NUMBER_PATTERN.fullmatch(value.strip()))
    if "udyam" in t:
        return bool(UDYAM_PATTERN.fullmatch(value.strip()))
    return bool(value.strip())


def apply_replacements(chunk: str, replacements: list[dict[str, str]]) -> str:
    updated = chunk
    # Replace longest strings first to reduce partial overlaps.
    sorted_replacements = sorted(
        replacements,
        key=lambda r: len(r.get("original_value", "")),
        reverse=True,
    )
    for item in sorted_replacements:
        src = item.get("original_value", "")
        dst = item.get("anonymized_value", "")
        if not src or src == dst:
            continue
        updated = re.sub(re.escape(src), dst, updated)
    return updated


def _chunk_report(
    chunk_index: int,
    original: str,
    anonymized: str,
    ner_groups: list[tuple[str, list[PiiDetection]]],
    combined_pii: list[PiiDetection],
    replacements: list[dict[str, str]],
    dropped_findings: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """Build a serializable report dict for one chunk (for UI/report JSON). ner_groups: (source_name, pii_list)."""
    def _pii_list(items: list[PiiDetection]) -> list[dict]:
        return [{"value": p.text, "pii_type": p.label, "score": round(p.score, 4)} for p in items]

    findings_with_source: list[dict] = []
    for p in combined_pii:
        found_by: list[str] = [name for name, pii in ner_groups if _group_contains_detection(pii, p.text, p.label)]
        if not found_by:
            found_by.append("audit")
        findings_with_source.append({
            "value": p.text,
            "pii_type": p.label,
            "score": round(p.score, 4),
            "found_by": found_by,
        })

    # Group duplicate/semi-duplicate findings (same or overlapping text + same label) as one; union found_by
    findings_deduped = _dedupe_findings(findings_with_source)

    report = {
        "chunk_index": chunk_index,
        "original": original,
        "anonymized": anonymized,
        "findings": findings_deduped,
        "replacements": list(replacements),
        "dropped_findings": list(dropped_findings) if dropped_findings else [],
    }
    for name, pii in ner_groups:
        report[name] = _pii_list(pii)
    # Backward compat for run_meta: presidio_count, gliner_count, qwen_count
    report.setdefault("presidio", [])
    report.setdefault("qwen", [])
    gliner_combined = []
    for name, pii in ner_groups:
        if name in (NER_NAME_XLARGE, NER_NAME_GRETELAI, NER_NAME_URCHADE, NER_NAME_ARABIC):
            gliner_combined.extend(_pii_list(pii))
    report["gliner"] = gliner_combined
    return report


def process_chunk(
    chunk: str,
    chunk_index: int,
    gliner_threshold: float,
    presidio_threshold: float,
    use_qwen_ner: bool,
    qwen_ner_threshold: float,
    min_ner_agreement: int,
    min_agreement_for_names: int,
    qwen_python: str,
    qwen_script: Path,
    progress_file: Path | None = None,
    file_name: str = "",
    total_chunks: int = 1,
    qwen_ner_model_name: str = "",
) -> tuple[str, bool, dict[str, Any]]:
    chunk_len = len(chunk)
    # Language detection via Qwen (Ollama): English -> all NERs except gliner_arabic; Arabic -> all except Presidio and gliner_xlarge
    lang = "en"
    if _detect_language is not None and use_qwen_ner:
        try:
            lang = _detect_language(chunk)
        except Exception as exc:  # pragma: no cover
            logger.warning("Language detection failed, assuming English: %s", exc)
    logger.info("Chunk %s: detected language=%s, NER requests (length=%s chars)", chunk_index, lang, chunk_len)
    _report_progress(progress_file, stage="language_detection", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, language=lang, chunk_size=chunk_len)

    # UI stage names for each GLiNER (must match armor.html data-stage)
    _gliner_stage_map = {
        NER_NAME_XLARGE: "gliner_xlarge",
        NER_NAME_GRETELAI: "gliner_gretelai",
        NER_NAME_URCHADE: "gliner_urchade",
        NER_NAME_ARABIC: "gliner_arabic",
    }

    ner_groups: list[tuple[str, list[PiiDetection]]] = []

    if lang == "ar":
        # Arabic only: gretelai, urchade, gliner_arabic (no Presidio, no xlarge)
        for model_id, name in [
            (GLINER_GRETELAI_ID, NER_NAME_GRETELAI),
            (GLINER_URCHADE_ID, NER_NAME_URCHADE),
            (GLINER_ARABIC_ID, NER_NAME_ARABIC),
        ]:
            ui_stage = _gliner_stage_map.get(name, "gliner")
            _report_progress(progress_file, stage=ui_stage, file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, gliner_model=name)
            logger.info("Requesting GLiNER %s (chunk length=%s chars)", name, chunk_len)
            model = _get_gliner_model(model_id)
            pii_list = detect_pii_with_gliner(model, chunk, threshold=gliner_threshold)
            ner_groups.append((name, pii_list))
        _report_progress(progress_file, stage="gliner_arabic", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, gliner_count=sum(len(p) for _, p in ner_groups))
    else:
        # All other languages (en, zh, etc.): Presidio + xlarge, gretelai, urchade — no Arabic GLiNER
        _report_progress(progress_file, stage="presidio", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, chunk_size=chunk_len)
        logger.info("Requesting Presidio (chunk length=%s chars)", chunk_len)
        presidio_pii = detect_pii_with_presidio(chunk, threshold=presidio_threshold)
        ner_groups.append(("presidio", presidio_pii))
        _report_progress(progress_file, stage="presidio", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, presidio_count=len(presidio_pii))
        for model_id, name in [
            (GLINER_XLARGE_ID, NER_NAME_XLARGE),
            (GLINER_GRETELAI_ID, NER_NAME_GRETELAI),
            (GLINER_URCHADE_ID, NER_NAME_URCHADE),
        ]:
            ui_stage = _gliner_stage_map.get(name, "gliner")
            _report_progress(progress_file, stage=ui_stage, file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, gliner_model=name)
            logger.info("Requesting GLiNER %s (chunk length=%s chars)", name, chunk_len)
            model = _get_gliner_model(model_id)
            pii_list = detect_pii_with_gliner(model, chunk, threshold=gliner_threshold)
            ner_groups.append((name, pii_list))
        _report_progress(progress_file, stage="gliner_urchade", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, gliner_count=sum(len(p) for _, p in ner_groups if _ != "presidio"))

    qwen_pii: list[PiiDetection] = []
    if use_qwen_ner and _qwen_ner_detect is not None:
        _report_progress(progress_file, stage="qwen_ner", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks)
        logger.info("Requesting Qwen NER (Ollama %s) (chunk length=%s chars)", qwen_ner_model_name or "?", chunk_len)
        try:
            qwen_raw = _qwen_ner_detect(chunk, threshold=qwen_ner_threshold)
            qwen_pii = [
                PiiDetection(text=x["text"], label=x["label"], score=float(x.get("score", 0.8)))
                for x in qwen_raw
                if isinstance(x, dict) and x.get("text") and x.get("label")
            ]
        except Exception as exc:  # pragma: no cover
            logger.warning("Qwen NER (Ollama) failed for chunk: %s", exc)
    qwen_pii = _normalize_pii_types_by_pattern(qwen_pii)
    ner_groups.append(("qwen", qwen_pii))
    _report_progress(progress_file, stage="qwen_ner", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, qwen_count=len(qwen_pii))

    detection_groups = [pii for _, pii in ner_groups]
    combined_pii = pii_ensemble_agreement(
        detection_groups,
        min_agreement=min_ner_agreement,
        min_agreement_for_names=min_agreement_for_names,
    )
    union = merge_pii_detections(detection_groups)
    agreed_set = {(p.text.lower(), p.label.lower()) for p in combined_pii}
    dropped_findings: list[dict[str, Any]] = []
    for p in union:
        key = (p.text.lower(), p.label.lower())
        if key in agreed_set:
            continue
        found_by = [name for name, pii in ner_groups if _group_contains_detection(pii, p.text, p.label)]
        dropped_findings.append({
            "value": p.text,
            "pii_type": p.label,
            "score": round(p.score, 4),
            "found_by": found_by,
            "reason": "agreement < required (kept only when at least {} detector(s) agree)".format(
                min_agreement_for_names if _is_person_name_label(p.label) else min_ner_agreement,
            ),
        })
    if dropped_findings:
        logger.info(
            "PII dropped (agreement): %s",
            json.dumps([{"value": d["value"], "pii_type": d["pii_type"], "found_by": d["found_by"]} for d in dropped_findings], ensure_ascii=True),
        )
    indian_pii = deterministic_audit_indian_ids(chunk)
    detected_values = {d.text for d in combined_pii}
    for p in indian_pii:
        if p.text not in detected_values:
            combined_pii.append(p)
            detected_values.add(p.text)
    audit_hits = regex_entropy_audit(chunk)
    for hit in audit_hits:
        if hit not in detected_values and hit.strip().lower() not in COMMON_NON_PII_WORDS:
            combined_pii.append(PiiDetection(text=hit, label="suspicious_token", score=1.0))
    combined_pii = _dedup_detections(combined_pii)
    combined_pii = _precision_filter(combined_pii)
    _report_progress(progress_file, stage="agreement", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, agreed_count=len(combined_pii))

    if not combined_pii:
        logger.info("no pii in the chunk")
        _report_progress(progress_file, stage="chunk_done", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, replacements_count=0)
        report = _chunk_report(chunk_index, chunk, chunk, ner_groups, [], [], dropped_findings=dropped_findings)
        return chunk, True, report

    _report_progress(progress_file, stage="anonymisation", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks)
    logger.info(
        "NER result counts (no chunk content): %s",
        ", ".join("%s=%s" % (name, len(pii)) for name, pii in ner_groups),
    )
    for name, pii in ner_groups:
        logger.info(
            "Piis found in the chunk by %s : %s",
            name, json.dumps([{"value": p.text, "pii_type": p.label} for p in pii], ensure_ascii=True),
        )
    logger.info(
        "Piis found in the chunk (min_agreement=%s) : %s",
        min_ner_agreement,
        json.dumps([{"value": p.text, "pii_type": p.label} for p in combined_pii], ensure_ascii=True),
    )
    r_system, r_user = build_qwen_replacement_prompt(chunk, combined_pii, lang=lang)
    qwen_response = call_qwen_json(qwen_python, qwen_script, r_system, r_user)
    replacements_raw = qwen_response.get("replacements", [])

    normalized: list[dict[str, str]] = []
    seen_values = set()
    for item in replacements_raw if isinstance(replacements_raw, list) else []:
        original = str(item.get("original_value", "")).strip()
        anonymized = str(item.get("anonymized_value", "")).strip()
        pii_type = str(item.get("pii_type", "")).strip() or "unknown"
        if not original:
            continue
        if original not in chunk:
            continue
        if "suspicious_token" in pii_type.lower():
            anonymized = partial_mask(original, visible_ratio=0.3)
        elif not datatype_match(anonymized, pii_type):
            anonymized = synthetic_value_for_type(pii_type, original, lang=lang)
        key = (original.lower(), pii_type.lower())
        if key in seen_values:
            continue
        seen_values.add(key)
        normalized.append(
            {
                "original_value": original,
                "anonymized_value": anonymized,
                "pii_type": pii_type,
            },
        )

    # Ensure every detected hit gets a replacement even if Qwen misses some.
    existing_originals = {r["original_value"].lower() for r in normalized}
    for p in combined_pii:
        if p.text.lower() in existing_originals:
            continue
        normalized.append(
            {
                "original_value": p.text,
                "anonymized_value": synthetic_value_for_type(p.label, p.text, lang=lang),
                "pii_type": p.label,
            },
        )

    logger.info("original and anonymized value in the chunk : %s", json.dumps(normalized, ensure_ascii=True))
    replaced_chunk = apply_replacements(chunk, normalized)
    logger.info("replacement done")
    _report_progress(progress_file, stage="chunk_done", file=file_name, chunk_index=chunk_index, total_chunks=total_chunks, replacements_count=len(normalized))
    report = _chunk_report(
        chunk_index, chunk, replaced_chunk,
        ner_groups, combined_pii, normalized,
        dropped_findings=dropped_findings,
    )
    return replaced_chunk, bool(normalized), report


def process_file(
    file_path: Path,
    args: argparse.Namespace,
) -> tuple[str, int, int, list[dict[str, Any]]]:
    logger.info("Processing file name start : %s", file_path.name)
    progress_file = getattr(args, "progress_file", None)
    _report_progress(progress_file, stage="extract", file=file_path.name)
    text = extract_text(file_path)
    chunks = chunk_text(text, chunk_size=args.chunk_size, chunk_overlap=args.chunk_overlap)
    _report_progress(progress_file, stage="chunking", file=file_path.name, total_chunks=len(chunks))

    anonymized_chunks: list[str] = []
    anonymized_count = 0
    not_anonymized_count = 0
    chunk_reports: list[dict[str, Any]] = []
    # Process only the first chunk (limit of 1 chunk)
    if chunks:
        index, chunk = 1, chunks[0]
        anonymized_chunk, success, report = process_chunk(
            chunk=chunk,
            chunk_index=index,
            gliner_threshold=args.gliner_threshold,
            presidio_threshold=args.presidio_threshold,
            use_qwen_ner=args.use_qwen_ner,
            qwen_ner_threshold=args.qwen_ner_threshold,
            min_ner_agreement=args.min_ner_agreement,
            min_agreement_for_names=getattr(args, "min_agreement_names", 1),
            qwen_python=args.qwen_python,
            qwen_script=args.qwen_script,
            progress_file=progress_file,
            file_name=file_path.name,
            total_chunks=len(chunks),
            qwen_ner_model_name=getattr(args, "qwen_ner_model", ""),
        )
        anonymized_chunks.append(anonymized_chunk)
        chunk_reports.append(report)
        if success:
            anonymized_count += 1
        else:
            not_anonymized_count += 1
    logger.info("Processing file name end : %s", file_path.name)
    return "\n".join(anonymized_chunks), anonymized_count, not_anonymized_count, chunk_reports


def main() -> None:
    args = parse_args()
    setup_logging()
    if AnalyzerEngine is not None:
        logger.info("Presidio: available (in-process)")
    else:
        logger.info(
            "Presidio: not available. To enable: pip install presidio-analyzer spacy && python -m spacy download en_core_web_sm",
        )
    logger.info(
        "Qwen NER (Ollama): %s",
        "enabled" if (args.use_qwen_ner and _qwen_ner_detect is not None) else "disabled or unavailable",
    )
    args.input_dir = args.input_dir.resolve()
    args.output_dir = args.output_dir.resolve()
    args.qwen_script = args.qwen_script.resolve()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    if not args.input_dir.is_dir():
        raise NotADirectoryError(f"Input dir not found: {args.input_dir}")
    if not args.qwen_script.is_file():
        raise FileNotFoundError(f"Qwen script not found: {args.qwen_script}")

    candidates = sorted(
        [
            path
            for path in args.input_dir.iterdir()
            if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES
        ],
        key=lambda p: p.name.lower(),
    )
    if getattr(args, "files", None) and args.files.strip():
        want = {s.strip() for s in args.files.split(",") if s.strip()}
        selected_files = [p for p in candidates if p.name in want]
        missing = want - {p.name for p in selected_files}
        if missing:
            logger.warning("Requested files not found in input-dir: %s", missing)
    else:
        selected_files = candidates[: args.num_files]
    if not selected_files:
        raise RuntimeError(f"No supported files found in {args.input_dir}")

    # GLiNER models are lazy-loaded per chunk (English: xlarge, gretelai, urchade; Arabic: gretelai, urchade, arabic)
    try:
        from qwen_ollama_ner_module import OLLAMA_NER_MODEL as _qwen_ner_model
    except Exception:
        _qwen_ner_model = os.environ.get("OLLAMA_NER_MODEL", "?")
    args.qwen_ner_model = getattr(args, "qwen_ner_model", _qwen_ner_model)
    logger.info(
        "NER: Presidio (EN only), GLiNER (xlarge/gretelai/urchade/arabic by language), Qwen NER (Ollama %s). Language detected per chunk via Qwen.",
        args.qwen_ner_model,
    )

    script_dir = Path(__file__).resolve().parent
    config_file = script_dir / "config.json"
    report_dir = args.report_dir.resolve()
    report_dir.mkdir(parents=True, exist_ok=True)

    _report_progress(args.progress_file, stage="starting", total_files=len(selected_files))
    total_anonymized = 0
    total_not_anonymized = 0
    file_reports: list[dict[str, Any]] = []
    for file_path in selected_files:
        anonymized_text, anonymized_count, not_anonymized_count, chunk_reports = process_file(
            file_path=file_path,
            args=args,
        )
        total_anonymized += anonymized_count
        total_not_anonymized += not_anonymized_count
        out_file = args.output_dir / f"{file_path.name}.anonymized.txt"
        out_file.write_text(anonymized_text, encoding="utf-8")

        original_text = extract_text(file_path)
        all_findings_raw: list[dict] = []
        all_replacements: list[dict] = []
        all_dropped_findings: list[dict] = []
        for cr in chunk_reports:
            all_findings_raw.extend(cr.get("findings", []))
            all_replacements.extend(cr.get("replacements", []))
            all_dropped_findings.extend(cr.get("dropped_findings", []))
        # Dedupe across chunks so same value (e.g. email in chunk 1 and 2) appears once with merged found_by
        all_findings = _dedupe_findings(all_findings_raw)
        file_reports.append({
            "file_name": file_path.name,
            "original_text": original_text,
            "anonymized_text": anonymized_text,
            "chunks": chunk_reports,
            "all_findings": all_findings,
            "all_replacements": all_replacements,
            "all_dropped_findings": all_dropped_findings,
            "chunks_anonymized": anonymized_count,
            "chunks_not_anonymized": not_anonymized_count,
        })

    logger.info("number of times all pii anonymized : %s", total_anonymized)
    logger.info("number of times it was not : %s", total_not_anonymized)

    if not args.no_report and file_reports:
        created_at = datetime.now().isoformat()
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        report_filename = f"{timestamp}_report.json"
        report_path = report_dir / report_filename
        report_payload = {
            "created_at": created_at,
            "source_files": [r["file_name"] for r in file_reports],
            "total_chunks_anonymized": total_anonymized,
            "total_chunks_not_anonymized": total_not_anonymized,
            "files": file_reports,
        }
        report_path.write_text(json.dumps(report_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        logger.info("Report written: %s", report_path)
        try:
            rel_report = report_path.relative_to(script_dir)
        except ValueError:
            rel_report = report_path
        config_data = {"latest_report": str(rel_report), "updated_at": created_at}
        config_file.write_text(json.dumps(config_data, indent=2), encoding="utf-8")
        logger.info("Config updated: %s (latest_report=%s)", config_file, rel_report)
        run_id = timestamp
        runs_dir = script_dir / "db" / "runs"
        runs_dir.mkdir(parents=True, exist_ok=True)
        run_dir = runs_dir / run_id
        run_dir.mkdir(parents=True, exist_ok=True)
        run_meta = {
            "run_id": run_id,
            "created_at": created_at,
            "report_path": str(rel_report),
            "total_chunks_anonymized": total_anonymized,
            "total_chunks_not_anonymized": total_not_anonymized,
            "files": [
                {
                    "file_name": r["file_name"],
                    "entity_count": len(r.get("all_findings", [])),
                    "entity_types": sorted(set(t for t in (f.get("pii_type", "") for f in r.get("all_findings", [])) if t)),
                    "chunks_processed": len(r.get("chunks", [])),
                    "chunk_logs": [
                        {
                            "chunk_index": c.get("chunk_index"),
                            "chunk_size": len(c.get("original", "")),
                            "presidio_count": len(c.get("presidio", [])),
                            "gliner_count": len(c.get("gliner", [])),
                            "qwen_count": len(c.get("qwen", [])),
                            "agreed_count": len(c.get("findings", [])),
                            "replacements_count": len(c.get("replacements", [])),
                        }
                        for c in r.get("chunks", [])
                    ],
                }
                for r in file_reports
            ],
        }
        (run_dir / "run_meta.json").write_text(json.dumps(run_meta, ensure_ascii=False, indent=2), encoding="utf-8")
        shutil.copy(report_path, run_dir / "report.json")
        logger.info("Run meta written: %s", run_dir / "run_meta.json")


if __name__ == "__main__":
    main()
